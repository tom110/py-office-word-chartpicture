import xlrd
import os
import wx
from docx import Document
from docx.shared import Cm


def get_path(wildcard):
    app=wx.App(None)
    style = wx.FD_OPEN | wx.FD_FILE_MUST_EXIST
    dialog = wx.FileDialog(None, '请选择Excel信息文件', wildcard=wildcard, style=style)
    if dialog.ShowModal() == wx.ID_OK:
        path = dialog.GetPath()
    else:
        path = None
    dialog.Destroy()
    return path

def get_dir(info):
    app=wx.App(None)
    style = wx.DD_DEFAULT_STYLE | wx.DD_NEW_DIR_BUTTON
    dialog = wx.DirDialog(None, info, style=style)
    if dialog.ShowModal() == wx.ID_OK:
        dir_path = dialog.GetPath()
    else:
        dir_path=None
    dialog.Destroy()
    return dir_path


desktop=os.path.join(os.path.expanduser("~"), 'Desktop')
excel_path=get_path('*.xlsx')
jpg_dir=get_dir("请选择图片所在文件夹")
doc_dir=get_dir("请选择word文档所在文件夹")


data = xlrd.open_workbook(excel_path)
table = data.sheets()[0]          #通过索引顺序获取
docs=table.col_values(0)
jpgs=table.col_values(1)

if not os.path.exists(desktop+'\\处理成果'):
    os.mkdir(desktop+'\\处理成果')

for d in docs:
    word = Document(doc_dir+'\\'+d)
    tables = word.tables
    p = tables[0].rows[20].cells[1].paragraphs[0]
    r = p.add_run()
    i= docs.index(d)
    r.add_picture(jpg_dir+'\\'+jpgs[i], width=Cm(4.95), height=Cm(3.5))
    # r.add_picture(jpg_dir + '\\' + jpgs[i])
    word.save(desktop+'\\处理成果\\'+d)









